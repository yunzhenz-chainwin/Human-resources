import base64
import binascii
import hashlib
import hmac
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.config import get_settings
from app.parsers import select_adapter
from app.services.ocr import extract_pdf_with_ocr

PARSER_VERSION = "adapters-2.2"

# OCR (Tesseract) 對中文常見的雜訊:字與字之間插入空白、標點變成小型/直排異體字。
# 只針對 OCR 產出的文字做正規化,數位 PDF 的文字層不受影響。
_OCR_PUNCT_MAP = {
    "﹕": ":", "︰": ":", "∶": ":",
    "﹒": ".",
    "﹣": "-", "－": "-", "‐": "-", "–": "-", "—": "-",
    "﹢": "+", "＋": "+",
    "＠": "@",
}
_CJK = r"一-鿿"
# 當空白至少一側是中文字時移除它(避免拆掉英文詞之間的正常空白)。
_OCR_SPACE_RE = re.compile(
    rf"(?<=[{_CJK}])[ \t]+(?=[{_CJK}0-9A-Za-z:.@+\-])"
    rf"|(?<=[0-9A-Za-z:.@+\-])[ \t]+(?=[{_CJK}])"
)


def normalize_ocr_text(text: str) -> str:
    """整理 OCR 文字:還原異體標點、移除中文字間被插入的空白。"""
    for src, dst in _OCR_PUNCT_MAP.items():
        text = text.replace(src, dst)
    previous = None
    while previous != text:
        previous = text
        text = _OCR_SPACE_RE.sub("", text)
    return text


@dataclass
class ParserResult:
    source_platform: str
    status: str
    raw_text: str
    payload: dict
    confidence: dict[str, float]
    overall_confidence: float
    source_confidence: float = 0.0
    source_evidence: list[dict] | None = None
    source_review_required: bool = True
    error_message: str | None = None


def _verify_structured_signature(payload_b64: str, signature_b64: str) -> bool:
    """Verify an HMAC-SHA256 signature over the base64 payload body.

    The signing key is the server ``auth_secret_key``.  A missing/short secret
    or any decode error fails closed so unsigned or tampered blobs are never
    trusted.
    """
    secret = get_settings().auth_secret_key.encode()
    if len(secret) < 32:
        return False
    try:
        provided = base64.urlsafe_b64decode(signature_b64 + "=" * (-len(signature_b64) % 4))
    except (binascii.Error, ValueError):
        return False
    expected = hmac.new(secret, payload_b64.encode("ascii"), hashlib.sha256).digest()
    return hmac.compare_digest(expected, provided)


def _structured_pdf_payload(path: Path) -> tuple[dict, bool] | None:
    """Read TalentHub fields embedded by the local PDF template.

    Returns the extracted payload together with ``verified`` — whether the blob
    carried a valid HMAC signature over its base64 body.  The ``THR1:`` prefix
    alone proves nothing about origin, so callers MUST treat ``verified=False``
    as untrusted and force manual review; an unsigned or forged blob is never
    granted the trusted fast-path.
    """
    try:
        metadata = PdfReader(str(path)).metadata
        subject = str(metadata.get("/Subject", "")) if metadata else ""
        if not subject.startswith("THR1:") or len(subject) > 20_000:
            return None
        payload_b64, _, signature_b64 = subject[5:].partition(".")
        verified = bool(signature_b64) and _verify_structured_signature(
            payload_b64, signature_b64
        )
        decoded = base64.b64decode(payload_b64, validate=True)
        source = json.loads(decoded.decode("utf-8"))
        if not isinstance(source, dict) or source.get("schema") != "talenthub.resume.v1":
            return None

        limits = {
            "name": 100,
            "email": 255,
            "phone": 50,
            "city": 50,
            "current_title": 100,
            "current_company": 100,
            "highest_education": 20,
            "expected_title": 100,
        }
        payload: dict = {}
        for field, limit in limits.items():
            value = source.get(field)
            if isinstance(value, str) and value.strip():
                payload[field] = value.strip()[:limit]
        years = source.get("total_years")
        if isinstance(years, (int, float)) and 0 <= float(years) <= 80:
            payload["total_years"] = float(years)
        for field, max_items in (("skills", 100), ("expected_cities", 30)):
            values = source.get(field)
            if isinstance(values, list):
                payload[field] = [
                    value.strip()[:100]
                    for value in values[:max_items]
                    if isinstance(value, str) and value.strip()
                ]
        return (payload, verified) if payload.get("name") else None
    except (ValueError, TypeError, json.JSONDecodeError, binascii.Error, PdfReadError):
        return None


def _structured_text(payload: dict) -> str:
    labels = (
        ("name", "姓名"),
        ("email", "Email"),
        ("phone", "電話"),
        ("city", "通訊地址"),
        ("current_title", "職務類別"),
        ("total_years", "工作年資"),
        ("highest_education", "學歷"),
        ("expected_title", "希望職務"),
    )
    lines = ["【TalentHub 結構化履歷】"]
    for field, label_name in labels:
        value = payload.get(field)
        if value is not None:
            suffix = " 年" if field == "total_years" else ""
            lines.append(f"{label_name}：{value}{suffix}")
    for field, label_name in (("skills", "技能與專長"), ("expected_cities", "希望工作地點")):
        values = payload.get(field) or []
        if values:
            lines.append(f"{label_name}：{'、'.join(values)}")
    return "\n".join(lines)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        result = extract_pdf_with_ocr(path)
        if result.needs_review:
            raise ValueError(result.error_message or "PDF requires manual review")
        return result.text
    if suffix == ".docx":
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend(cell.text for row in table.rows for cell in row.cells)
        return "\n".join(paragraphs)
    if suffix == ".doc":
        raise ValueError(
            "Legacy DOC files are stored safely but require manual review or conversion"
        )
    raise ValueError("Unsupported resume format")


def detect_platform(text: str, requested: str) -> str:
    """Compatibility entry point used by older callers and tests."""
    return select_adapter(text, requested).platform


def parse_text(text: str, requested_platform: str = "generic") -> ParserResult:
    text = text.replace("\x00", "").strip()
    parsed = select_adapter(text, requested_platform)
    populated = [score for score in parsed.confidence.values() if score > 0]
    overall = (
        round(sum(populated) / len(parsed.confidence), 2) if populated else 0.0
    )
    has_identity = bool(
        parsed.payload.get("name")
        and (parsed.payload.get("email") or parsed.payload.get("phone"))
    )
    status = (
        "parsed"
        if parsed.layout_recognized and has_identity and overall >= 0.55
        else "needs_review"
    )
    if parsed.source_review_required:
        status = "needs_review"
    error = None
    if not parsed.layout_recognized:
        error = (
            f"Unknown {parsed.platform} layout; parsed with {parsed.version} "
            "and requires manual review"
        )
    return ParserResult(
        parsed.platform,
        status,
        text,
        parsed.payload,
        parsed.confidence,
        overall,
        parsed.source_confidence,
        parsed.source_evidence,
        parsed.source_review_required,
        error,
    )


def parse_resume(path: Path, requested_platform: str) -> ParserResult:
    if path.suffix.lower() == ".pdf":
        structured = _structured_pdf_payload(path)
        if structured is not None:
            payload, verified = structured
            # A forged THR1 prefix must not bypass review: only a cryptographically
            # verified blob is granted full trust; otherwise the fields are still
            # surfaced but flagged for manual review at reduced confidence.
            confidence = {
                field: (1.0 if verified else 0.75)
                for field, value in payload.items()
                if value is not None
            }
            has_identity = bool(
                payload.get("name")
                and (payload.get("email") or payload.get("phone"))
            )
            return ParserResult(
                "direct",
                "parsed" if (verified and has_identity) else "needs_review",
                _structured_text(payload),
                payload,
                confidence,
                1.0 if verified else 0.75,
                1.0 if verified else 0.0,
                [
                    {
                        "type": (
                            "structured_pdf_metadata"
                            if verified
                            else "unverified_structured_pdf_metadata"
                        ),
                        "marker": "talenthub.resume.v1",
                        "weight": 1.0 if verified else 0.0,
                    }
                ],
                not verified,
                None,
            )
        ocr_result = extract_pdf_with_ocr(path)
        if ocr_result.needs_review:
            return ParserResult(
                requested_platform,
                "needs_review",
                "",
                {},
                {},
                0.0,
                0.0,
                [{"type": "ocr_failure", "marker": "no_text", "weight": 0.0}],
                True,
                ocr_result.error_message or "PDF requires manual review",
            )
        text = ocr_result.text
        is_ocr = ocr_result.status == "ocr_extracted"
        if is_ocr:
            text = normalize_ocr_text(text)
        result = parse_text(text, requested_platform)
        if is_ocr:
            # Regex certainty is not OCR certainty. Scanned fields always need
            # human review even when their recognized shape looks valid.
            result.confidence = {
                field: round(score * 0.78, 2)
                for field, score in result.confidence.items()
            }
            populated = [score for score in result.confidence.values() if score > 0]
            result.overall_confidence = (
                round(sum(populated) / len(result.confidence), 2) if populated else 0.0
            )
            result.status = "needs_review"
            result.source_evidence = list(result.source_evidence or []) + [
                {"type": "ocr_extracted", "marker": ocr_result.provider, "weight": 0.0}
            ]
        return result
    try:
        return parse_text(extract_text(path), requested_platform)
    except Exception as exc:
        return ParserResult(
            requested_platform,
            "needs_review" if path.suffix.lower() == ".doc" else "failed",
            "",
            {},
            {},
            0.0,
            0.0,
            [{"type": "extraction_failure", "marker": path.suffix.lower(), "weight": 0.0}],
            True,
            str(exc)[:1000],
        )
