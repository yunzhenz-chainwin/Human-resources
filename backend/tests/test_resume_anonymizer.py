import importlib.util
import io
import subprocess
import sys
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from pypdf import PdfWriter

ROOT = Path(__file__).resolve().parents[2]
MODULE_PATH = ROOT / "resume_anonymizer.py"
SPEC = importlib.util.spec_from_file_location("standalone_resume_anonymizer", MODULE_PATH)
assert SPEC is not None and SPEC.loader is not None
resume_anonymizer = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = resume_anonymizer
SPEC.loader.exec_module(resume_anonymizer)


def _blank_pdf(page_count: int = 1) -> bytes:
    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=612, height=792)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def test_scanned_pdf_fails_closed_with_private_ocr_guidance() -> None:
    with pytest.raises(resume_anonymizer.ScannedPDFError) as captured:
        resume_anonymizer.extract_pdf(_blank_pdf())

    message = str(captured.value)
    assert "掃描型 PDF" in message
    assert "不會直接產出去識別結果" in message
    assert "公司核准的本機 OCR" in message
    assert "人工確認" in message
    assert "未核准的第三方 OCR" in message


def test_pdf_limits_are_checked_before_anonymization() -> None:
    with pytest.raises(ValueError, match="超過 20 頁上限"):
        resume_anonymizer.extract_pdf(_blank_pdf(21))

    with pytest.raises(ValueError, match="超過 10 MB 上限"):
        resume_anonymizer.extract_pdf(b"x" * (resume_anonymizer.PDF_MAX_BYTES + 1))


def test_cli_pdf_input_uses_the_same_scanned_pdf_guard() -> None:
    data = _blank_pdf()

    class ScannedPath:
        suffix = ".pdf"

        @staticmethod
        def stat():
            return type("Stat", (), {"st_size": len(data)})()

        @staticmethod
        def read_bytes() -> bytes:
            return data

    with pytest.raises(resume_anonymizer.ScannedPDFError):
        resume_anonymizer.read_resume_input(ScannedPath())


def test_selectable_pdf_text_is_returned(monkeypatch: pytest.MonkeyPatch) -> None:
    expected = "姓名：王小明\n工作經驗：後端工程師五年，負責金融系統 API 開發與維護"

    class FakePage:
        def extract_text(self) -> str:
            return expected

    class FakeReader:
        is_encrypted = False
        pages = [FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", lambda _stream: FakeReader())
    assert resume_anonymizer.extract_pdf(b"%PDF-test") == expected


def test_scanned_pdf_error_reports_page_text_counts(monkeypatch: pytest.MonkeyPatch) -> None:
    class Page:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    class Reader:
        is_encrypted = False
        pages = [Page("姓名"), Page(""), Page("電話 0912")]

    monkeypatch.setattr("pypdf.PdfReader", lambda _stream: Reader())
    with pytest.raises(resume_anonymizer.ScannedPDFError) as captured:
        resume_anonymizer.extract_pdf(b"%PDF-test")

    error = captured.value
    assert error.page_count == 3
    assert error.page_characters == (2, 0, 6)
    assert error.threshold == 20
    message = str(error)
    assert "PDF 共 3 頁" in message
    assert "目前擷取 8 字，至少需要 20 字" in message
    assert "第1頁 2 字" in message
    assert "第2頁 0 字" in message
    assert "第3頁 6 字" in message
    assert "可能原因" in message


def test_web_page_escapes_resume_and_error_content() -> None:
    handler = object.__new__(resume_anonymizer.WebHandler)
    page = handler._page(
        source="</textarea><script>alert('pii')</script>",
        result="<b>[NAME]</b>",
        error="<script>error</script>",
    ).decode("utf-8")

    assert "<script>" not in page
    assert "&lt;/textarea&gt;&lt;script&gt;" in page
    assert "&lt;b&gt;[NAME]&lt;/b&gt;" in page


def test_docx_input_extracts_paragraphs_and_table_cells() -> None:
    document = Document()
    document.add_paragraph("姓名 王小名")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Email"
    table.cell(0, 1).text = "person@example.com"
    output = io.BytesIO()
    document.save(output)

    text = resume_anonymizer.extract_docx(output.getvalue())
    anonymized, summary = resume_anonymizer.anonymize(text)
    assert "姓名 [NAME]" in anonymized
    assert "[EMAIL]" in anonymized
    assert summary.replacements["name_label"] == 1


def test_text_input_accepts_utf8_and_rejects_unknown_suffix() -> None:
    class TextPath:
        suffix = ".txt"

        @staticmethod
        def read_bytes() -> bytes:
            return "姓名 王小名\n電話 0912-345-678".encode()

    assert "姓名 王小名" in resume_anonymizer.read_resume_input(TextPath())

    class UnknownPath:
        suffix = ".rtf"

    unknown = UnknownPath()
    with pytest.raises(ValueError, match="不支援的檔案格式"):
        resume_anonymizer.read_resume_input(unknown)
def test_ocr_quality_flags_low_confidence() -> None:
    assert resume_anonymizer._ocr_quality([95.0, 90.0], "name candidate with enough text") >= 0.90
    assert resume_anonymizer._ocr_quality([30.0, 40.0], "name candidate with enough text") < 0.75


def test_ocr_resolution_is_300_dpi() -> None:
    assert resume_anonymizer.OCR_DPI == 300


def test_local_ocr_cleans_raw_resume_when_poppler_fails(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    scratch = ROOT / "backend"
    token = f"cleanup-test-{uuid4().hex}"
    expected_pdf = scratch / f".resume-ocr-{token}.pdf"
    monkeypatch.chdir(scratch)
    monkeypatch.setattr(resume_anonymizer.shutil, "which", lambda name: name)
    monkeypatch.setattr(resume_anonymizer.secrets, "token_hex", lambda _length: token)

    calls = 0

    def fake_run(args, **_kwargs):
        nonlocal calls
        calls += 1
        if calls == 1:
            return type("Completed", (), {"stdout": "List of available languages:\nchi_tra\n"})()
        raise subprocess.CalledProcessError(1, args)

    monkeypatch.setattr(resume_anonymizer.subprocess, "run", fake_run)

    with pytest.raises(RuntimeError, match="Poppler"):
        resume_anonymizer._local_ocr_pdf(b"private resume", 1)

    assert not expected_pdf.exists()
