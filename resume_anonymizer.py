#!/usr/bin/env python3
"""Standalone resume de-identification utility.

The module deliberately has no dependency on the TalentHub backend.  It can be
used from the command line or as a small local web page (``--web``).
"""

from __future__ import annotations

import argparse
import cgi
import html
import io
import json
import os
import re
import secrets
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import datetime
from dataclasses import asdict, dataclass
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, urlparse

PDF_MAX_BYTES = 10 * 1024 * 1024
PDF_MAX_PAGES = 20
PDF_MAX_TEXT_CHARACTERS = 250_000
OCR_TIMEOUT_SECONDS = 90
OCR_DPI = 200
WEB_MAX_REQUEST_BYTES = PDF_MAX_BYTES + 1 * 1024 * 1024
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".text"}
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ScannedPDFError(ValueError):
    """A PDF has no trustworthy selectable text.

    The attributes are intentionally exposed so the web UI/CLI can explain
    which pages failed extraction instead of showing only a generic warning.
    """

    def __init__(
        self,
        message: str,
        *,
        page_count: int = 0,
        page_characters: tuple[int, ...] = (),
        threshold: int = 20,
    ) -> None:
        super().__init__(message)
        self.page_count = page_count
        self.page_characters = page_characters
        self.threshold = threshold


@dataclass
class Summary:
    input_characters: int
    output_characters: int
    replacements: dict[str, int]


PATTERNS: tuple[tuple[str, re.Pattern[str], str], ...] = (
    ("email", re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I), "[EMAIL]"),
    ("phone", re.compile(r"(?<!\d)(?:\+?886[-\s]?|0)9\d{2}[-\s]?\d{3}[-\s]?\d{3}(?!\d)"), "[PHONE]"),
    ("taiwan_id", re.compile(r"(?<![A-Z0-9])[A-Z][12]\d{8}(?![A-Z0-9])", re.I), "[ID]"),
)


def anonymize(text: str) -> tuple[str, Summary]:
    """Replace common personal identifiers while retaining layout."""
    result = text
    counts: dict[str, int] = {}
    for name, pattern, replacement in PATTERNS:
        result, counts[name] = pattern.subn(replacement, result)

    # Label-aware rules are deliberately conservative; this avoids redacting
    # company/project names that merely happen to look like a person's name.
    name_label = re.compile(
        r"(?im)^(\s*(?:姓名|名字|求職者姓名|應徵者姓名|name)\s*(?:[:：]|\s+)\s*)[^\r\n]+"
    )
    result, counts["name_label"] = name_label.subn(r"\1[NAME]", result)
    address_label = re.compile(
        r"(?im)^(\s*(?:地址|住址|居住地區|所在地|location|address)\s*(?:[:：]|\s+)\s*)[^\r\n]+"
    )
    result, counts["address_label"] = address_label.subn(r"\1[ADDRESS]", result)
    # Common Taiwan address forms when no label is present.
    address = re.compile(
        r"(?<![\w\u4e00-\u9fff])(?:臺|台)[北中南東西]市[^\r\n,，]{2,40}(?:區|鄉|鎮|市)[^\r\n,，]{0,40}"
    )
    result, counts["address"] = address.subn("[ADDRESS]", result)
    return result, Summary(len(text), len(result), counts)


def _local_ocr_pdf(data: bytes, page_count: int) -> tuple[str, tuple[int, ...]]:
    """Run local Poppler + Tesseract only, with bounded resources."""
    tesseract, pdftoppm = shutil.which("tesseract"), shutil.which("pdftoppm")
    if not tesseract or not pdftoppm:
        raise RuntimeError("找不到本機 OCR：請安裝 Tesseract OCR 與 Poppler，並將工具加入 PATH。")
    started = time.monotonic()
    # Use unique filenames in the local workspace. Some managed Windows
    # profiles deny subprocess access to redirected temp directories.
    token = secrets.token_hex(12)
    class _WorkspaceScratch:
        def __enter__(self):
            return str(Path.cwd())
        def __exit__(self, *_args):
            return False
    with _WorkspaceScratch() as work:
        pdf = Path(work) / f".resume-ocr-{token}.pdf"; prefix = Path(work) / f".resume-ocr-{token}-page"; pdf.write_bytes(data)
        try:
            subprocess.run([pdftoppm, "-f", "1", "-l", str(page_count), "-r", str(OCR_DPI), "-png", str(pdf), str(prefix)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=OCR_TIMEOUT_SECONDS, shell=False)
        except subprocess.TimeoutExpired as exc:
            raise TimeoutError(f"本機 OCR 超過 {OCR_TIMEOUT_SECONDS} 秒，已停止處理。") from exc
        except (OSError, subprocess.CalledProcessError) as exc:
            raise RuntimeError("Poppler 無法將 PDF 轉成影像，請確認 PDF 未損壞。") from exc
        parts, counts = [], []
        for index in range(1, page_count + 1):
            image = Path(f"{prefix}-{index}.png")
            if not image.exists(): parts.append(""); counts.append(0); continue
            try:
                out = subprocess.run([tesseract, str(image), "stdout", "-l", "chi_tra+eng", "--psm", "6"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=max(5, OCR_TIMEOUT_SECONDS - int(time.monotonic() - started)), shell=False)
            except subprocess.TimeoutExpired as exc:
                raise TimeoutError(f"本機 OCR 超過 {OCR_TIMEOUT_SECONDS} 秒，已停止處理。") from exc
            except (OSError, subprocess.CalledProcessError) as exc:
                raise RuntimeError("Tesseract 無法辨識頁面，請確認已安裝 chi_tra 語言包。") from exc
            value = out.stdout.decode("utf-8", errors="replace").strip(); parts.append(value); counts.append(len(re.sub(r"\s+", "", value)))
        result = "\n".join(parts).strip()
        # Best-effort cleanup; OCR output is never retained in the result.
        for candidate in [pdf, *(Path(f"{prefix}-{i}.png") for i in range(1, page_count + 1))]:
            try:
                candidate.unlink(missing_ok=True)
            except OSError:
                pass
        return result, tuple(counts)


def extract_pdf(data: bytes) -> str:
    if len(data) > PDF_MAX_BYTES:
        raise ValueError(f"PDF 檔案超過 {PDF_MAX_BYTES // (1024 * 1024)} MB 上限")
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on environment
        raise RuntimeError("PDF 解析需要 pypdf，請執行 python -m pip install pypdf") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise ValueError("PDF 受密碼保護，無法安全讀取")
        if len(reader.pages) > PDF_MAX_PAGES:
            raise ValueError(f"PDF 頁數超過 {PDF_MAX_PAGES} 頁上限")
        page_texts = [page.extract_text() or "" for page in reader.pages]
        page_characters = tuple(len(re.sub(r"\s+", "", value)) for value in page_texts)
        text = "\n".join(page_texts).strip()
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("PDF 無法解析，請確認檔案未損壞") from exc
    extracted_characters = len(re.sub(r"\s+", "", text))
    if extracted_characters < 20:
        try:
            ocr_text, _ocr_pages = _local_ocr_pdf(data, len(reader.pages))
        except (RuntimeError, TimeoutError) as exc:
            details = "、".join(f"第{i}頁 {count} 字" for i, count in enumerate(page_characters, 1))
            raise ScannedPDFError(f"PDF 共 {len(reader.pages)} 頁，目前擷取 {extracted_characters} 字，至少需要 20 字。各頁文字量：{details}。可能原因：掃描圖片、空白頁或文字層損壞。本機 OCR 未完成：{exc}。請使用本機 OCR 後重新上傳，或貼上人工確認文字。", page_count=len(reader.pages), page_characters=page_characters, threshold=20) from exc
        if len(re.sub(r"\s+", "", ocr_text)) >= 20:
            return ocr_text
        page_details = "、".join(
            f"第{index}頁 {count} 字" for index, count in enumerate(page_characters, 1)
        ) or "無頁面文字"
        raise ScannedPDFError(
            f"偵測到掃描型 PDF 或文字層不足：PDF 共 {len(reader.pages)} 頁，"
            f"目前擷取 {extracted_characters} 字，至少需要 {20} 字。各頁文字量：{page_details}。"
            "可能原因：PDF 是掃描圖片、文字層損壞、頁面為空白或文字無法選取。"
            "請使用公司核准的本機 OCR 重新產生可搜尋 PDF，確認每頁文字後再上傳，並由人工確認；"
            "請勿使用未核准的第三方 OCR。本工具不會直接產出去識別結果或送出履歷內容。",
            page_count=len(reader.pages),
            page_characters=page_characters,
            threshold=20,
        )
    if len(text) > PDF_MAX_TEXT_CHARACTERS:
        raise ValueError(f"PDF 文字不可超過 {PDF_MAX_TEXT_CHARACTERS} 字")
    return text


def extract_docx(data: bytes) -> str:
    """Extract paragraphs and table cells from a DOCX document."""
    if len(data) > PDF_MAX_BYTES:
        raise ValueError(f"DOCX 檔案不可超過 {PDF_MAX_BYTES // (1024 * 1024)} MB")
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("DOCX 解析需要 python-docx，請執行 python -m pip install python-docx") from exc
    try:
        document = Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise ValueError("DOCX 無法解析，請確認檔案未損壞") from exc
    if len(text) > PDF_MAX_TEXT_CHARACTERS:
        raise ValueError(f"DOCX 文字不可超過 {PDF_MAX_TEXT_CHARACTERS} 字")
    return text


def render_docx(text: str) -> bytes:
    """Create a sanitized DOCX for direct backend import."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("DOCX export requires python-docx") from exc
    document = Document()
    for line in text.splitlines() or [""]:
        document.add_paragraph(line)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def extract_text(data: bytes, suffix: str = ".txt") -> str:
    """Decode plain text, accepting UTF-8 (with BOM), UTF-16, or legacy fallback."""
    if len(data) > PDF_MAX_BYTES:
        raise ValueError(f"文字檔不可超過 {PDF_MAX_BYTES // (1024 * 1024)} MB")
    for encoding in ("utf-8-sig", "utf-16", "cp950"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError("文字檔編碼無法辨識，請轉成 UTF-8 後再試")
    if len(text) > PDF_MAX_TEXT_CHARACTERS:
        raise ValueError(f"文字不可超過 {PDF_MAX_TEXT_CHARACTERS} 字")
    return text


def read_resume_input(path: Path) -> str:
    suffix = path.suffix.casefold()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError("不支援的檔案格式；請使用 PDF、DOCX 或 TXT")
    data = path.read_bytes()
    if suffix == ".pdf":
        return extract_pdf(data)
    if suffix == ".docx":
        return extract_docx(data)
    return extract_text(data, suffix)


def _extract_uploaded(filename: str, data: bytes) -> str:
    suffix = Path(filename or "").suffix.casefold()
    if suffix == ".pdf":
        return extract_pdf(data)
    if suffix == ".docx":
        return extract_docx(data)
    if suffix in {".txt", ".text", ""}:
        return extract_text(data, suffix)
    raise ValueError("不支援的檔案格式；請使用 PDF、DOCX 或 TXT")


def _run_local_ocr(data: bytes) -> tuple[str, int, str]:
    """Run the existing bounded local OCR provider for a scanned PDF."""
    backend_root = Path(__file__).resolve().parent / "backend"
    if str(backend_root) not in sys.path:
        sys.path.insert(0, str(backend_root))
    try:
        from app.services.ocr import OCRLimits, LocalTesseractProvider, extract_pdf_with_ocr
    except Exception as exc:
        raise RuntimeError("本機 OCR 元件尚未安裝，請確認 PyMuPDF 與 Tesseract") from exc
    temp_root = Path(tempfile.gettempdir()) / "talenthub-resume-anonymizer"
    temp_root.mkdir(parents=True, exist_ok=True)
    pdf_path = temp_root / f"ocr-{secrets.token_hex(12)}.pdf"
    pdf_path.write_bytes(data)
    try:
        provider = LocalTesseractProvider()
        if not provider.available():
            raise RuntimeError("找不到本機 Tesseract OCR 或 PyMuPDF，請先安裝公司核准的 OCR 元件")
        result = extract_pdf_with_ocr(pdf_path, provider=provider, limits=OCRLimits())
        if result.needs_review or not result.text.strip():
            raise RuntimeError(result.error_message or "本機 OCR 未辨識到可用文字")
        return result.text, result.page_count, result.provider
    finally:
        pdf_path.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="獨立履歷去識別化工具")
    parser.add_argument("input", nargs="?", default="-", help="輸入檔案；- 代表標準輸入")
    parser.add_argument("-o", "--output", help="輸出檔案")
    parser.add_argument("--summary", action="store_true", help="將摘要輸出到 stderr")
    parser.add_argument("--web", action="store_true", help="啟動本機網頁")
    parser.add_argument("--port", type=int, default=8765)
    args = parser.parse_args()
    if args.web:
        return serve_web(args.port)
    text = sys.stdin.read() if args.input == "-" else read_resume_input(Path(args.input))
    anonymized, summary = anonymize(text)
    if args.output:
        output_path = Path(args.output)
        if output_path.suffix.casefold() == ".docx":
            output_path.write_bytes(render_docx(anonymized))
        else:
            output_path.write_text(anonymized, encoding="utf-8")
    else:
        sys.stdout.write(anonymized)
    if args.summary:
        print(json.dumps(asdict(summary), ensure_ascii=False), file=sys.stderr)
    return 0


class WebHandler(BaseHTTPRequestHandler):
    _results: dict[str, tuple[str, bytes]] = {}
    _ocr_sources: dict[str, tuple[str, bytes]] = {}
    _used_names: set[str] = set()
    _result_dir = Path(tempfile.gettempdir()) / "talenthub-resume-anonymizer"

    @classmethod
    def _store_result(cls, token: str, filename: str, payload: bytes) -> None:
        cls._results[token] = (filename, payload)
        cls._result_dir.mkdir(parents=True, exist_ok=True)
        (cls._result_dir / f"{token}.name").write_text(filename, encoding="utf-8")
        (cls._result_dir / f"{token}.bin").write_bytes(payload)

    @classmethod
    def _get_result(cls, token: str) -> tuple[str, bytes] | None:
        item = cls._results.get(token)
        if item:
            return item
        if not re.fullmatch(r"[A-Za-z0-9_-]{12,80}", token):
            return None
        try:
            filename = (cls._result_dir / f"{token}.name").read_text(encoding="utf-8")
            try:
                payload = (cls._result_dir / f"{token}.bin").read_bytes()
            except FileNotFoundError:
                payload = (cls._result_dir / f"{token}.txt").read_bytes()
        except (FileNotFoundError, OSError, UnicodeError):
            return None
        cls._results[token] = (filename, payload)
        return filename, payload

    @classmethod
    def _store_ocr_source(cls, token: str, filename: str, payload: bytes) -> None:
        cls._ocr_sources[token] = (filename, payload)

    @classmethod
    def _get_ocr_source(cls, token: str) -> tuple[str, bytes] | None:
        return cls._ocr_sources.get(token)

    @classmethod
    def _filename(cls, original: str = "", suffix: str = ".txt") -> str:
        date = datetime.now().strftime("%Y%m%d")
        stem = Path(original).stem if original else "resume"
        stem = re.sub(r"[^A-Za-z0-9\u4e00-\u9fff._-]+", "_", stem).strip("._") or "resume"
        base = f"{stem}_{date}"
        suffix = suffix if suffix in {".txt", ".docx"} else ".txt"
        name = f"{base}{suffix}"
        index = 1
        while name in cls._used_names:
            name = f"{base}_{index:02d}{suffix}"
            index += 1
        cls._used_names.add(name)
        return name

    def _page(self, source: str = "", result: str = "", summary: Summary | None = None, error: str = "", download: str = "", filename: str = "", ocr_token: str = "", ocr_info: str = "") -> bytes:
        total = sum(summary.replacements.values()) if summary else 0
        rows = "".join(f"<li><span>{html.escape(k)}</span><b>{v}</b></li>" for k, v in (summary.replacements.items() if summary else []) if v)
        rows = rows or "<li><span>未偵測到可替換欄位</span><b>0</b></li>"
        summary_html = ""
        ocr_action = ""
        if ocr_token:
            ocr_action = ("<form method='post' class='ocr-form'><input type='hidden' name='ocr_token' value='" + html.escape(ocr_token) + "'>"
                          "<button type='submit' name='ocr_action' value='run'>使用本機 OCR</button></form>")
        ocr_info_html = ("<div class='ocr-info'><strong>本機 OCR 處理資訊</strong><p>" + html.escape(ocr_info) +
                         "</p><small>OCR 完成後已進入去識別化流程；請人工確認結果完整性。</small></div>") if ocr_info else ""
        if summary:
            summary_html = f"<section class='result-summary'><div class='stat'><small>原始字數</small><strong>{summary.input_characters:,}</strong></div><div class='stat'><small>處理後字數</small><strong>{summary.output_characters:,}</strong></div><div class='stat accent'><small>已替換欄位</small><strong>{total}</strong></div><div class='replacement-list'><h3>替換明細</h3><ul>{rows}</ul></div></section>"
        error_html = f"<div class='alert'><strong>無法安全處理這份檔案</strong><p>{html.escape(error)}</p><ol><li>掃描 PDF 請先用公司核准的本機 OCR 轉成可搜尋 PDF。</li><li>確認文字可以反白複製，或改上傳 DOCX／TXT。</li><li>也可以直接貼上人工確認過的履歷文字。</li></ol></div>" if error else ""
        download_html = f"<div class='download-box'><div><strong>處理完成，可下載去識別化檔案</strong><small>統一輸出為 UTF-8 純文字 TXT，方便人工覆核。</small></div><a class='download-link' href='/download/{html.escape(download)}'>下載 {html.escape(filename)}</a></div>" if download else ""
        document = f"""<!doctype html><html lang='zh-Hant'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width, initial-scale=1'><title>履歷去識別化｜TalentHub 工具</title><style>
:root{{--ink:#173b3a;--muted:#64817d;--teal:#087f70;--line:#d7e8e4;--soft:#eff8f5;--gold:#d89b28}}*{{box-sizing:border-box}}body{{margin:0;background:#f6faf9;color:var(--ink);font-family:Arial,'Microsoft JhengHei',sans-serif}}.shell{{max-width:1240px;margin:0 auto;padding:34px 26px 56px}}header{{display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:24px}}h1{{font-size:32px;margin:0 0 8px;color:#086b62;letter-spacing:.04em}}.subtitle,.hint{{color:var(--muted);font-size:13px}}.badge{{padding:8px 12px;border-radius:999px;background:#e4f3ef;color:var(--teal);font-size:13px;font-weight:700}}.notice{{margin:16px 0;padding:13px 16px;border:1px solid #f1dcad;border-left:4px solid var(--gold);border-radius:9px;background:#fffaf0;color:#6e5625;font-size:13px;line-height:1.6}}.alert{{margin:16px 0;padding:16px 18px;border:1px solid #efb8b1;border-left:4px solid #c9574e;border-radius:9px;background:#fff3f1;color:#7e302a;line-height:1.6}}.alert p{{margin:6px 0}}.alert ol{{margin:8px 0 0 20px;padding:0}}form{{margin-top:20px}}.grid{{display:grid;grid-template-columns:1fr 1fr;gap:20px}}.panel{{background:white;border:1px solid var(--line);border-radius:14px;padding:18px;box-shadow:0 5px 18px rgba(20,80,70,.05)}}.panel-head{{display:flex;justify-content:space-between;align-items:center;margin-bottom:10px}}label,.panel-title{{font-weight:700;font-size:15px}}input[type=file]{{width:100%;padding:10px;border:1px dashed #a7cbc2;border-radius:8px;background:#fbfefd;margin:8px 0 10px}}textarea{{width:100%;height:330px;resize:vertical;padding:14px;border:1px solid #c4ddd8;border-radius:9px;font:14px/1.65 Consolas,'Microsoft JhengHei',monospace;color:#203d3a;background:#fcfefe}}textarea:focus{{outline:2px solid #a7ded4;border-color:var(--teal)}}.actions{{display:flex;gap:10px;align-items:center;margin-top:16px}}button,.secondary{{border:0;border-radius:8px;padding:12px 20px;background:var(--teal);color:white;font-weight:700;cursor:pointer;text-decoration:none;font-size:14px}}.secondary{{background:white;color:var(--teal);border:1px solid #abd1c9}}.result-summary{{display:grid;grid-template-columns:repeat(3,1fr) 2fr;gap:10px;align-items:stretch;margin-top:20px}}.stat,.replacement-list{{background:white;border:1px solid var(--line);border-radius:12px;padding:14px}}.stat small{{display:block;color:var(--muted);font-size:12px;margin-bottom:8px}}.stat strong{{font-size:25px;color:var(--teal)}}.stat.accent{{background:#e8f6f1}}.replacement-list h3{{font-size:14px;margin:0 0 8px}}.replacement-list ul{{list-style:none;margin:0;padding:0;display:grid;grid-template-columns:1fr 1fr;gap:5px 16px;font-size:13px}}.replacement-list li{{display:flex;justify-content:space-between;border-bottom:1px solid #edf4f2;padding:3px 0}}.replacement-list b{{color:var(--teal)}}.download-box{{display:flex;justify-content:space-between;align-items:center;gap:16px;margin-top:16px;padding:15px 18px;background:#eaf7f3;border:1px solid #bfe0d8;border-radius:11px}}.download-box small{{display:block;color:var(--muted);font-weight:400;margin-top:4px}}.download-link{{padding:10px 14px;border-radius:8px;background:var(--teal);color:white;font-weight:700;text-decoration:none;white-space:nowrap}}footer{{margin-top:24px;color:var(--muted);font-size:12px;text-align:center}}@media(max-width:850px){{.shell{{padding:24px 14px}}header{{display:block}}.grid,.result-summary{{grid-template-columns:1fr}}.download-box{{display:block}}.download-link{{display:inline-block;margin-top:12px}}}}
</style></head><body><main class='shell'><header><div><h1>履歷去識別化</h1><div class='subtitle'>獨立 Python 工具 · 不連接 TalentHub 後台 · 僅在本機處理</div></div><span class='badge'>安全工作區</span></header><div class='notice'>支援 PDF、DOCX、TXT 與直接貼上的純文字。掃描 PDF 若沒有可搜尋文字，系統會停止處理並告訴你原因，不會冒險產出未完整遮蔽的結果。</div>{error_html}<form method='post' enctype='multipart/form-data'><div class='grid'><section class='panel'><div class='panel-head'><label>原始履歷</label><span class='hint'>可上傳檔案或貼上文字</span></div><input type='file' name='file' accept='.pdf,.docx,.txt,.text,application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document,text/plain'><div class='hint'>PDF 請確認文字可以反白複製；掃描檔請先使用公司核准的本機 OCR。</div><textarea name='text' placeholder='請貼上履歷文字'>{html.escape(source)}</textarea></section><section class='panel'><div class='panel-head'><span class='panel-title'>去識別化結果</span><span class='hint'>可供人工覆核</span></div><textarea readonly placeholder='處理後結果會顯示在這裡'>{html.escape(result)}</textarea>{download_html}</section></div><div class='actions'><button type='submit'>開始去識別化</button><a class='secondary' href='/?reset=1'>重新開始</a></div></form>{summary_html}<footer>請在匯入人才庫前再次人工確認姓名、地址、電話、Email 與身分證等個資是否已完整遮蔽。</footer></main></body></html>"""
        if ocr_action:
            document = document.replace("</body>", ocr_action + "</body>")
        if ocr_info_html:
            document = document.replace("</body>", ocr_info_html + "</body>")
        if filename.casefold().endswith(".docx"):
            document = document.replace("TXT", "DOCX")
        return document.encode("utf-8")

    def _send_page(self, status_code: int, body: bytes) -> None:
        self.send_response(status_code)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        # The page is a local, stateful tool.  Do not let the browser reuse a
        # previous error/result when the user refreshes or starts a new file.
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path
        if path.startswith("/download/"):
            item = self._get_result(path.rsplit("/", 1)[-1])
            if not item:
                self._send_page(404, self._page(error="下載連結已失效，請重新處理履歷。"))
                return
            filename, payload = item
            self.send_response(200)
            content_type = DOCX_MIME if filename.casefold().endswith(".docx") else "text/plain; charset=utf-8"
            self.send_header("Content-Type", content_type)
            self.send_header("Cache-Control", "no-store")
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(filename)}")
            self.send_header("Content-Length", str(len(payload)))
            self.end_headers()
            self.wfile.write(payload)
            return
        query = parse_qs(parsed.query)
        if query.get("reset", [""])[0] == "1":
            self.send_response(303)
            self.send_header("Location", "/")
            self.send_header("Cache-Control", "no-store")
            self.end_headers()
            return
        token = query.get("token", [""])[0]
        item = self._get_result(token)
        if item:
            filename, payload = item
            if filename.casefold().endswith(".docx"):
                display_result = extract_docx(payload)
            else:
                display_result = payload.decode("utf-8")
            ocr_info = ""
            if query.get("ocr", [""])[0] == "1":
                pages = query.get("pages", ["?"])[0]
                chars = query.get("chars", ["?"])[0]
                ocr_info = f"已使用本機 OCR；頁數：{pages}；擷取字數：{chars}；引擎：Tesseract。"
            self._send_page(200, self._page(result=display_result, download=token, filename=filename, ocr_info=ocr_info))
        else:
            self._send_page(200, self._page())

    def do_POST(self) -> None:
        try:
            length = int(self.headers.get("Content-Length", "0"))
            if length < 0 or length > WEB_MAX_REQUEST_BYTES:
                raise ValueError(f"請求不可超過 {WEB_MAX_REQUEST_BYTES // (1024 * 1024)} MB")
            raw = self.rfile.read(length)
            content_type = self.headers.get("Content-Type", "")
            if content_type.startswith("application/x-www-form-urlencoded"):
                fields = parse_qs(raw.decode("utf-8"), keep_blank_values=True)
                ocr_token = fields.get("ocr_token", [""])[0]
                source_item = self._get_ocr_source(ocr_token)
                if fields.get("ocr_action", [""])[0] == "run" and source_item:
                    original_name, pdf_data = source_item
                    try:
                        ocr_text, page_count, provider = _run_local_ocr(pdf_data)
                        result, summary = anonymize(ocr_text)
                        token = secrets.token_urlsafe(18)
                        filename = self._filename(original_name, ".txt")
                        self._store_result(token, filename, result.encode("utf-8"))
                        self.send_response(303)
                        self.send_header("Location", f"/?token={token}&ocr=1&pages={page_count}&chars={len(ocr_text)}")
                        self.end_headers()
                        return
                    except Exception as exc:
                        self._send_page(400, self._page(error=str(exc), ocr_token=ocr_token, ocr_info="OCR 未完成，請確認本機 OCR 元件與 PDF 品質"))
                        return
            source = ""
            original_name = ""
            original_suffix = ".txt"
            if content_type.startswith("multipart/form-data"):
                form = cgi.FieldStorage(fp=io.BytesIO(raw), headers=self.headers, environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": content_type})
                source = form.getfirst("text", "")
                upload = form["file"] if "file" in form else (form["pdf"] if "pdf" in form else None)
                if upload is not None and getattr(upload, "filename", "") and getattr(upload, "file", None):
                    original_name = upload.filename
                    original_suffix = Path(original_name).suffix.casefold()
                    upload_data = upload.file.read()
                    try:
                        source = _extract_uploaded(upload.filename, upload_data)
                    except ScannedPDFError as exc:
                        ocr_token = secrets.token_urlsafe(18)
                        self._store_ocr_source(ocr_token, original_name, upload_data)
                        detail = f"PDF 共 {exc.page_count} 頁；擷取文字不足 {exc.threshold} 字。可按下「使用本機 OCR」自動辨識。"
                        self._send_page(400, self._page(error=detail, ocr_token=ocr_token, ocr_info="尚未執行 OCR"))
                        return
            else:
                source = parse_qs(raw.decode("utf-8"), keep_blank_values=True).get("text", [""])[0]
            result, summary = anonymize(source)
            token = secrets.token_urlsafe(18)
            output_suffix = ".docx" if original_suffix == ".docx" else ".txt"
            filename = self._filename(original_name, output_suffix)
            payload = render_docx(result) if output_suffix == ".docx" else result.encode("utf-8")
            self._store_result(token, filename, payload)
            self.send_response(303)
            self.send_header("Location", f"/?token={token}")
            self.end_headers()
            return
        except Exception as exc:
            self._send_page(400, self._page(error=str(exc)))

    def log_message(self, *_args: object) -> None:
        return


def serve_web(port: int) -> int:
    server = ThreadingHTTPServer(("127.0.0.1", port), WebHandler)
    server.daemon_threads = True
    print(f"履歷去識別化工具：http://127.0.0.1:{port}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        return 0
    finally:
        server.server_close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
